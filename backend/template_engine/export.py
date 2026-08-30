"""Template document export — Markdown (native), PDF, DOCX.

Pure functions: (document, definition) in, bytes/str out. No DB access,
so these are directly unit-testable.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, List

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

from .models import TemplateDefinition, TemplateDocument


def _field_value_lines(value: Any) -> List[str]:
    if isinstance(value, list):
        return [f"- {v}" for v in value] if value else ["_(vide)_"]
    if value in (None, ""):
        return ["_(vide)_"]
    return [str(value)]


def to_markdown(document: TemplateDocument, definition: TemplateDefinition) -> str:
    lines = [
        f"# {document.title}",
        "",
        f"_Template : {definition.title} (v{document.definition_version})_",
        "",
    ]
    for field in definition.fields:
        lines.append(f"## {field.label}")
        lines.extend(_field_value_lines(document.values.get(field.key)))
        lines.append("")
    return "\n".join(lines)


def to_pdf(document: TemplateDocument, definition: TemplateDefinition) -> bytes:
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 25 * mm

    c.setFont("Helvetica-Bold", 18)
    c.drawString(20 * mm, y, document.title)
    y -= 10 * mm
    c.setFont("Helvetica-Oblique", 10)
    c.drawString(
        20 * mm, y, f"Template : {definition.title} (v{document.definition_version})"
    )
    y -= 12 * mm

    for field in definition.fields:
        if y < 30 * mm:
            c.showPage()
            y = height - 25 * mm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(20 * mm, y, field.label)
        y -= 7 * mm
        c.setFont("Helvetica", 10)
        for line in _field_value_lines(document.values.get(field.key)):
            for wrapped in _wrap(line, 95):
                if y < 20 * mm:
                    c.showPage()
                    y = height - 25 * mm
                c.drawString(24 * mm, y, wrapped)
                y -= 5.5 * mm
        y -= 4 * mm

    c.showPage()
    c.save()
    return buf.getvalue()


def _wrap(text: str, width: int) -> List[str]:
    if len(text) <= width:
        return [text]
    words = text.split(" ")
    out: List[str] = []
    current = ""
    for w in words:
        if len(current) + len(w) + 1 > width:
            out.append(current)
            current = w
        else:
            current = f"{current} {w}".strip()
    if current:
        out.append(current)
    return out or [""]


def to_docx(document: TemplateDocument, definition: TemplateDefinition) -> bytes:
    doc = DocxDocument()
    doc.add_heading(document.title, level=1)
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        f"Template : {definition.title} (v{document.definition_version})"
    ).italic = True

    for field in definition.fields:
        doc.add_heading(field.label, level=2)
        value = document.values.get(field.key)
        if isinstance(value, list):
            if value:
                for item in value:
                    doc.add_paragraph(str(item), style="List Bullet")
            else:
                doc.add_paragraph("(vide)")
        else:
            doc.add_paragraph(str(value) if value not in (None, "") else "(vide)")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
